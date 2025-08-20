from docx import Document
from pathlib import Path
from typing import Union, cast
from auto_documentation.markdown_converter.html_validator import HtmlNode, SupportedTags
from bs4 import BeautifulSoup


class HtmlToWordConverter:
    def __init__(
        self,
        html_node: Union[HtmlNode, None],
        test_output_path: Union[Path, str, None],
        html_file_path: Union[Path, str, None] = None,
    ):
        self.html_file_path = html_file_path
        self.html_node = html_node
        self.test_output_path = test_output_path
        self.doc = Document()
        self.html_as_string = None

        if self.html_file_path and self.html_node is None:
            self.open_html_file()
        elif self.html_node:
            self.html_as_string = self.html_node.display_string()

        self.convert()
        # self.save_to_file(self.test_output_path)

    def open_html_file(self):
        if not self.html_file_path.absolute():
            raise ValueError("HTML file path must be absolute")
        if not self.html_file_path.exists():
            raise ValueError("HTML file path must exist")
        if not self.html_file_path.is_file():
            raise ValueError("HTML file path must be a file")
        if self.html_file_path.suffix != ".html":
            raise ValueError("HTML file path must have a .html extension")

        raw_data = self.html_file_path.read_text()
        self.html_as_string = raw_data

    def recursive_convert(self, node: BeautifulSoup, parent_element: None):
        for child in node.contents:
            # Will need to also process siblings here
            name = child.name
            match name:
                case (SupportedTags.HEADER, SupportedTags.BODY):
                    continue
                case (SupportedTags.BR, SupportedTags.HR):
                    self.doc.add_page_break()
                case (
                    SupportedTags.P,
                    SupportedTags.DIV,
                ):
                    paragraph = self.doc.add_paragraph()
                    text = child.content
                    paragraph.add_run(text)

                case SupportedTags.TABLE:
                    table = self.doc.add_table()
                case SupportedTags.TH:
                    ...
                case SupportedTags.TR:
                    ...
                case SupportedTags.LI:
                    ...
                case SupportedTags.UL:
                    ...
                case (
                    SupportedTags.H1,
                    SupportedTags.H2,
                    SupportedTags.H3,
                    SupportedTags.H4,
                    SupportedTags.H5,
                    SupportedTags.H6,
                ):
                    ...
                case SupportedTags.A:
                    ...
                case SupportedTags.IMG:
                    # Content here would have to be bytes
                    # self.doc.add_picture(child.content)
                    ...
                case (
                    SupportedTags.EM,
                    SupportedTags.SPAN,
                    SupportedTags.STRONG,
                ):
                    parent_element = cast(Document.Paragraph, parent_element)
                    parent_element

            if child.children:
                self.recursive_convert(child, parent_element=parent_element)

    def convert(self):
        try:
            as_soup = BeautifulSoup(self.html_as_string, "html.parser")
            print(as_soup)
            return True
        except Exception as e:
            raise e

    def get_doc(self):
        return self.doc

    def convert_to_bytes(self): ...

    def save_to_file(self, file_path: Union[Path, str]):
        self.doc.save(file_path)
